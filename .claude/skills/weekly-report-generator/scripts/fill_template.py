#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word 模板填充脚本

作者：老王
日期：2026-01-17
"""

import sys
import json
import re
import argparse
from pathlib import Path
from typing import Dict, Any

try:
    from .common import fix_windows_console_encoding
except ImportError:
    from common import fix_windows_console_encoding

fix_windows_console_encoding()


def fill_word_template(
    template_path: str,
    data_path: str,
    output_path: str,
    filename: str
) -> str:
    """填充 Word 模板"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "需要安装 python-docx 库：pip install python-docx"
        )

    print(f"📖 读取模板文件：{template_path}")
    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"无法读取 Word 模板：{e}")

    print(f"📖 读取数据文件：{data_path}")
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        raise ValueError(f"无法读取 JSON 数据：{e}")

    content_map = _extract_content_map(data)
    print(f"📊 找到 {len(content_map)} 个章节内容")

    print("✍️  正在填充模板...")
    _fill_document(doc, content_map)

    output_dir = Path(output_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_file = output_dir / filename
    print(f"💾 保存文件：{output_file}")
    doc.save(output_file)

    return str(output_file.resolve())


def _extract_content_map(data: Dict[str, Any]) -> Dict[str, str]:
    """从 JSON 数据中提取章节内容映射

    支持格式：
    1. 嵌套结构：{"sections": [{"title": "...", "content": "..."}]}
    2. 扁平结构：{"本周工作情况": "...", "下周工作计划": "..."}
    3. report_content结构：{"report_content": {"本周工作情况": [...], "下周工作计划": [...]}}
    """
    def _convert_to_string(value) -> str:
        """将各种类型转换为字符串"""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, list):
            return '\n'.join(str(item).strip() for item in value if item)
        if isinstance(value, dict):
            return _convert_to_string(value.get("content", ""))
        return str(value).strip()

    def _add_content(key: str, value: Any) -> None:
        """添加非空内容到映射表"""
        if value:
            content = _convert_to_string(value)
            if content:
                content_map[key] = content

    content_map = {}
    special_keys = {"title", "sections", "report_content", "week", "start_date", "end_date"}

    # 格式1：从 data["sections"] 中提取
    for section in data.get("sections", []):
        title = section.get("title", "").strip()
        if title:
            _add_content(title, section.get("content", ""))

    # 格式2：从 data["report_content"] 中提取
    for key, value in data.get("report_content", {}).items():
        _add_content(key, value)

    # 格式3：直接从顶层字段提取
    for key, value in data.items():
        if key not in special_keys:
            _add_content(key, value)

    return content_map


def _normalize_title(title: str) -> str:
    """标准化标题，用于智能匹配（移除末尾标点、多余空格）

    优先进行精确匹配，如果失败则尝试标准化匹配
    """
    # 移除常见的末尾标点符号
    title = re.sub(r'[：:。，、\s]+$', '', title)
    title = title.strip()
    return title


def _find_matching_title(template_title: str, content_map: Dict[str, str]) -> str:
    """在 content_map 中查找匹配的标题（精确匹配 / 标准化匹配）"""
    if template_title in content_map:
        return template_title

    normalized_template = _normalize_title(template_title)
    for key in content_map.keys():
        if _normalize_title(key) == normalized_template:
            return key

    return ""


def _fill_document(doc, content_map: Dict[str, str]) -> None:
    """填充 Word 文档"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    filled_sections = set()

    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 检测标题（跳过 Heading 1，它通常是文档总标题）
        if para.style.name.startswith("Heading") and para.style.name != "Heading 1":
            current_section_title = text

            matched_key = _find_matching_title(current_section_title, content_map)

            if matched_key and current_section_title not in filled_sections:
                content = content_map[matched_key]
                content_lines = [line.strip() for line in content.split('\n') if line.strip()]

                if not content_lines:
                    continue

                # 在标题后按正确顺序插入内容段落
                # 从最后一行开始倒序插入，这样顺序才是正确的
                for line in reversed(content_lines):
                    new_para = para.insert_paragraph_before(line)
                    # 将新段落移动到标题后
                    para._element.addnext(new_para._element)

                filled_sections.add(current_section_title)


def _is_placeholder(text: str) -> bool:
    """判断文本是否是占位符"""
    placeholder_patterns = [
        "{{", "}}", "【", "】", "content", "内容", "请填写",
        "填写", "此处", "输入", "描述"
    ]

    text_lower = text.lower()
    return any(pattern.lower() in text_lower for pattern in placeholder_patterns)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="使用 JSON 数据填充 Word 模板"
    )
    parser.add_argument(
        "--template",
        type=str,
        required=True,
        help="Word 模板文件路径（.docx）"
    )
    parser.add_argument(
        "--data",
        type=str,
        required=True,
        help="JSON 数据文件路径"
    )
    parser.add_argument(
        "--output",
        type=str,
        required=True,
        help="输出完整路径（包括文件名）"
    )
    parser.add_argument(
        "--title",
        type=str,
        default="",
        help="文档标题（可选）"
    )

    args = parser.parse_args()

    print("=" * 60)
    print("✍️  Word 模板填充脚本")
    print("=" * 60)
    print()

    try:
        output_path = Path(args.output)
        output_dir = output_path.parent
        output_filename = output_path.name

        output_file = fill_word_template(
            template_path=args.template,
            data_path=args.data,
            output_path=str(output_dir),
            filename=output_filename
        )

        print()
        print("=" * 60)
        print(f"✅ 模板填充完成！")
        print(f"📄 输出文件：{output_file}")
        print("=" * 60)

    except Exception as e:
        print()
        print("=" * 60)
        print(f"❌ 填充失败：{e}")
        print("=" * 60)
        sys.exit(1)


if __name__ == "__main__":
    main()
