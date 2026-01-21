#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""模板分析脚本

作者：老王
日期：2026-01-16
"""

import sys
import json
import re
import argparse
from pathlib import Path
from typing import Dict, List, Optional, Any

try:
    from .common import fix_windows_console_encoding, extract_variables
except ImportError:
    from common import fix_windows_console_encoding, extract_variables

fix_windows_console_encoding()


class WordAnalyzer:
    """Word 文档分析器"""

    def __init__(self):
        self.elements = []

    def analyze(self, doc_path: str) -> Dict[str, Any]:
        """分析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            return {
                "type": "word",
                "error": "需要安装 python-docx 库：pip install python-docx",
                "raw_content": ""
            }

        try:
            doc = Document(doc_path)
        except Exception as e:
            return {
                "type": "word",
                "error": f"无法读取 Word 文件：{e}",
                "raw_content": ""
            }

        content = "\n".join([para.text for para in doc.paragraphs])
        title = self._extract_title(doc)
        sections = self._extract_sections(doc)
        variables = self._extract_variables(content)

        return {
            "type": "word",
            "structure": {
                "title": title,
                "sections": sections
            },
            "variables": variables,
            "raw_content": content
        }

    def _extract_title(self, doc) -> str:
        """提取文档标题"""
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading") and para.text.strip():
                return para.text.strip()

        for para in doc.paragraphs:
            if para.text.strip():
                return para.text.strip()

        return ""

    def _extract_sections(self, doc) -> List[Dict[str, Any]]:
        """提取章节结构"""
        sections = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = 1
                if para.style.name != "Heading":
                    try:
                        level = int(para.style.name.replace("Heading ", ""))
                    except ValueError:
                        level = 1

                sections.append({
                    "level": level,
                    "title": para.text.strip(),
                    "content": "",
                    "style_name": para.style.name
                })

        return sections

    def _extract_variables(self, content: str) -> Dict[str, Dict[str, Any]]:
        """提取变量占位符（使用公共模块）"""
        return extract_variables(content)


def analyze_markdown_template(template_path: str) -> dict:
    """分析 Markdown 模板文件"""
    template_path = Path(template_path)

    try:
        with open(template_path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        return {
            "type": "markdown",
            "error": f"无法读取文件：{e}",
            "raw_content": ""
        }

    title = ""
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()

    sections = []
    section_pattern = r'^(#{1,6})\s+(.+?)$'

    current_section = None
    for line in content.split('\n'):
        match = re.match(section_pattern, line)
        if match:
            level = len(match.group(1))
            title_section = match.group(2).strip()

            section = {
                "level": level,
                "title": title_section,
                "content": ""
            }
            sections.append(section)
            current_section = section
        elif current_section:
            if current_section["content"]:
                current_section["content"] += "\n" + line
            else:
                current_section["content"] = line

    required_vars = ["起始日期", "结束日期", "工作内容列表"]
    variables = extract_variables(content, required_vars)

    return {
        "type": "markdown",
        "structure": {
            "title": title,
            "sections": sections
        },
        "variables": variables,
        "raw_content": content
    }


def analyze_template(template_path: str) -> dict:
    """分析模板文件（自动检测文件类型）"""
    template_path = Path(template_path)

    if not template_path.exists():
        return {
            "type": "unknown",
            "error": f"文件不存在：{template_path}",
            "raw_content": ""
        }

    suffix = template_path.suffix.lower()

    if suffix == ".md":
        return analyze_markdown_template(template_path)
    elif suffix == ".docx":
        analyzer = WordAnalyzer()
        return analyzer.analyze(template_path)
    else:
        return {
            "type": "unknown",
            "error": f"不支持的文件类型：{suffix}，请使用 .md 或 .docx 文件",
            "raw_content": ""
        }


def main():
    """主函数"""
    print("=" * 60)
    print("🔍 模板分析脚本启动")
    print("=" * 60)
    print()

    parser = argparse.ArgumentParser(
        description="分析周报模板文件，提取结构和变量"
    )
    parser.add_argument(
        "--template",
        type=str,
        required=True,
        help="模板文件路径（支持 .md、.docx）"
    )
    parser.add_argument(
        "--output",
        type=str,
        help="输出 JSON 文件路径（可选，不提供则打印到控制台）"
    )

    args = parser.parse_args()

    print("📋 输入参数：")
    print(f"  - 模板文件：{args.template}")
    print(f"  - 输出文件：{args.output if args.output else '（控制台输出）'}")
    print()

    template_path = Path(args.template)
    if not template_path.exists():
        print(f"❌ 错误：文件不存在：{args.template}")
        sys.exit(1)

    suffix = template_path.suffix.lower()
    print(f"📄 检测文件类型：{suffix}")
    print()

    if suffix not in [".md", ".docx"]:
        print(f"❌ 错误：不支持的文件类型：{suffix}")
        print("💡 支持的文件类型：.md、.docx")
        sys.exit(1)

    print("⏳ 正在分析模板...")
    result = analyze_template(args.template)

    if "error" in result:
        print(f"❌ 分析失败：{result['error']}")
        sys.exit(1)

    print("✅ 模板分析成功")
    print()

    print("📊 分析结果摘要：")
    print(f"  - 文件类型：{result['type']}")
    if result['type'] != 'unknown':
        structure = result.get('structure', {})
        title = structure.get('title', '（无标题）')
        sections = structure.get('sections', [])
        variables = result.get('variables', {})

        print(f"  - 文档标题：{title}")
        print(f"  - 章节数量：{len(sections)}")
        print(f"  - 变量数量：{len(variables)}")

        if sections:
            print("  - 章节列表：")
            for section in sections[:5]:
                indent = "  " * (section['level'])
                print(f"    {indent}- {section['title']}")
            if len(sections) > 5:
                print(f"    ... 还有 {len(sections) - 5} 个章节")

        if variables:
            print("  - 变量列表：")
            for i, (placeholder, var_info) in enumerate(variables.items(), 1):
                required = "必需" if var_info.get('required') else "可选"
                print(f"    {i}. {placeholder} ({required})")
    print()

    print("💾 正在输出 JSON...")
    json_str = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        print(f"   写入文件：{args.output}")
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(json_str)
        print(f"✅ 分析结果已保存到：{output_path}")
        print(f"   文件大小：{len(json_str)} 字节")
    else:
        print("   输出到控制台")
        print()
        print("-" * 60)
        print(json_str)
        print("-" * 60)

    print()
    print("=" * 60)
    print("✅ 模板分析完成")
    print("=" * 60)


if __name__ == "__main__":
    main()
