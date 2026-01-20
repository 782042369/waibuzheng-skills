#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 模板填充脚本

功能：
- 读取 Word 模板文件（.docx）
- 读取 JSON 数据文件（包含要填充的内容）
- 替换模板中的变量占位符（{{变量名}}）
- 保持原始格式、样式、布局
- 输出填充后的 Word 文件

作者：老王
日期：2026-01-17
"""

import sys
import json
import re
import argparse
from pathlib import Path
from typing import Dict, Any

# 导入公共模块
try:
    from .common import fix_windows_console_encoding
except ImportError:
    from common import fix_windows_console_encoding

# 修复 Windows 控制台编码问题
fix_windows_console_encoding()


def fill_word_template(
    template_path: str,
    data_path: str,
    output_path: str,
    filename: str
) -> str:
    """
    填充 Word 模板

    Args:
        template_path: Word 模板文件路径
        data_path: JSON 数据文件路径
        output_path: 输出目录路径
        filename: 输出文件名

    Returns:
        输出文件的完整路径
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "需要安装 python-docx 库：pip install python-docx"
        )

    # 读取模板文件
    print(f"📖 读取模板文件：{template_path}")
    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"无法读取 Word 模板：{e}")

    # 读取数据文件
    print(f"📖 读取数据文件：{data_path}")
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        raise ValueError(f"无法读取 JSON 数据：{e}")

    # 提取章节内容映射
    content_map = _extract_content_map(data)
    print(f"📊 找到 {len(content_map)} 个章节内容")

    # 填充模板
    print("✍️  正在填充模板...")
    _fill_document(doc, content_map)

    # 确保输出目录存在
    output_dir = Path(output_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 保存文件
    output_file = output_dir / filename
    print(f"💾 保存文件：{output_file}")
    doc.save(output_file)

    return str(output_file.resolve())


def _extract_content_map(data: Dict[str, Any]) -> Dict[str, str]:
    """
    从 JSON 数据中提取章节内容映射

    Args:
        data: JSON 数据

    Returns:
        {章节标题: 章节内容} 的映射字典

    支持多种格式：
    1. 嵌套结构：{"sections": [{"title": "...", "content": "..."}]}
    2. 扁平结构：{"本周工作情况": "...", "下周工作计划": "..."}
    3. report_content结构：{"report_content": {"本周工作情况": [...]，"下周工作计划": [...]}}
    """

    def _convert_to_string(value) -> str:
        """将任意值转换为字符串"""
        if isinstance(value, str):
            return value.strip()
        elif isinstance(value, list):
            # 数组用换行符连接
            return '\n'.join(str(item).strip() for item in value if item)
        elif isinstance(value, dict):
            # 递归处理字典
            return _convert_to_string(value.get("content", ""))
        else:
            return str(value).strip()

    content_map = {}

    # 格式1：从 data["sections"] 中提取（嵌套结构）
    sections = data.get("sections", [])
    if sections:
        for section in sections:
            title = section.get("title", "").strip()
            raw_content = section.get("content", "")

            if title and raw_content:
                content = _convert_to_string(raw_content)
                if content:
                    content_map[title] = content

    # 格式2：从 data["report_content"] 中提取（report_content结构）
    report_content = data.get("report_content", {})
    if report_content:
        for key, value in report_content.items():
            if value:  # 确保值不为空
                content = _convert_to_string(value)
                if content:
                    content_map[key] = content

    # 格式3：直接从顶层字段提取（扁平结构）
    # 跳过特殊字段
    special_keys = {"title", "sections", "report_content", "week", "start_date", "end_date"}
    for key, value in data.items():
        if key not in special_keys and value:
            content = _convert_to_string(value)
            if content:
                content_map[key] = content

    return content_map


def _normalize_title(title: str) -> str:
    """
    标准化标题，用于智能匹配

    Args:
        title: 原始标题

    Returns:
        标准化后的标题

    处理规则：
    - 移除末尾的标点符号（：:。，）
    - 移除多余空格
    - 转为小写（用于匹配）
    """
    # 移除末尾的标点符号
    title = re.sub(r'[：:。，、\s]+$', '', title)
    # 移除开头和结尾的空格
    title = title.strip()
    return title


def _find_matching_title(template_title: str, content_map: Dict[str, str]) -> str:
    """
    在 content_map 中查找匹配的标题

    Args:
        template_title: 模板中的标题
        content_map: 内容映射字典

    Returns:
        匹配的key，如果未找到返回空字符串
    """
    # 1. 精确匹配
    if template_title in content_map:
        return template_title

    # 2. 标准化后匹配（忽略末尾标点）
    normalized_template = _normalize_title(template_title)
    for key in content_map.keys():
        if _normalize_title(key) == normalized_template:
            return key

    # 3. 未找到匹配
    return ""


def _fill_document(doc, content_map: Dict[str, str]) -> None:
    """
    填充 Word 文档

    Args:
        doc: docx.Document 对象
        content_map: {章节标题: 章节内容} 的映射字典
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    filled_sections = set()

    # 遍历所有段落（倒序，这样插入段落不会影响索引）
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 检测是否是标题（Heading 样式，排除 Heading 1）
        if para.style.name.startswith("Heading") and not para.style.name == "Heading 1":
            current_section_title = text

            # 智能查找匹配的标题
            matched_key = _find_matching_title(current_section_title, content_map)

            if matched_key and current_section_title not in filled_sections:
                content = content_map[matched_key]

                # 检查标题后面是否有空段落或占位符段落
                has_empty_paragraph = False
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_text = next_para.text.strip()
                    if not next_text or _is_placeholder(next_text):
                        # 下一个段落是空的或是占位符，可以直接填充
                        has_empty_paragraph = True
                        # 将内容按换行符分割成多个段落
                        content_lines = content.split('\n')
                        # 清空第一个段落并填充第一行内容
                        next_para.clear()
                        if content_lines and content_lines[0].strip():
                            next_para.add_run(content_lines[0].strip())
                        # 在后面插入剩余的段落
                        for j in range(1, len(content_lines)):
                            if content_lines[j].strip():
                                new_para = next_para.insert_paragraph_before(content_lines[j].strip())
                                next_para._p.addnext(new_para._p)

                # 如果没有空段落，在标题后插入新段落
                if not has_empty_paragraph:
                    # 将内容按换行符分割成多个段落
                    content_lines = content.split('\n')
                    # 倒序插入段落（这样顺序才正确）
                    for line in reversed(content_lines):
                        if line.strip():  # 跳过空行
                            new_para = para.insert_paragraph_before(line.strip())
                            # 将新段落移到标题后面
                            para._p.addnext(new_para._p)

                # 标记为已填充
                filled_sections.add(current_section_title)


def _is_placeholder(text: str) -> bool:
    """
    判断文本是否是占位符

    Args:
        text: 文本内容

    Returns:
        是否是占位符
    """
    placeholder_patterns = [
        "{{", "}}", "【", "】", "content", "内容", "请填写",
        "填写", "此处", "输入", "描述"
    ]

    text_lower = text.lower()
    return any(pattern.lower() in text_lower for pattern in placeholder_patterns)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="使用 JSON 数据填充 Word 模板"
    )
    parser.add_argument(
        "--template",
        type=str,
        required=True,
        help="Word 模板文件路径（.docx）"
    )
    parser.add_argument(
        "--data",
        type=str,
        required=True,
        help="JSON 数据文件路径"
    )
    parser.add_argument(
        "--output",
        type=str,
        required=True,
        help="输出完整路径（包括文件名）"
    )
    parser.add_argument(
        "--title",
        type=str,
        default="",
        help="文档标题（可选）"
    )

    args = parser.parse_args()

    print("=" * 60)
    print("✍️  Word 模板填充脚本")
    print("=" * 60)
    print()

    try:
        # 解析输出路径（分离目录和文件名）
        output_path = Path(args.output)
        output_dir = output_path.parent
        output_filename = output_path.name

        # 调用填充函数
        output_file = fill_word_template(
            template_path=args.template,
            data_path=args.data,
            output_path=str(output_dir),
            filename=output_filename
        )

        print()
        print("=" * 60)
        print(f"✅ 模板填充完成！")
        print(f"📄 输出文件：{output_file}")
        print("=" * 60)

    except Exception as e:
        print()
        print("=" * 60)
        print(f"❌ 填充失败：{e}")
        print("=" * 60)
        sys.exit(1)


if __name__ == "__main__":
    main()
