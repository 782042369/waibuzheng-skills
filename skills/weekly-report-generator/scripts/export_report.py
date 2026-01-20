#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
周报导出脚本（重构版）

功能：
- 将AI生成的完整周报内容写入文件
- 支持Markdown、Word两种格式
- 根据文件名扩展名自动检测输出格式
- 支持自动文件名生成
- 支持日期格式化（用于文件名）
- 使用专业的 Markdown → Word 转换器

作者：老王
日期：2026-01-16
"""

import sys
import argparse
import re
from pathlib import Path
from datetime import datetime
from typing import Optional

# 导入公共模块
try:
    from .common import fix_windows_console_encoding
except ImportError:
    from common import fix_windows_console_encoding

# 修复 Windows 控制台编码问题
fix_windows_console_encoding()


# ========== Markdown → Word 转换器（改进版） ==========

class MarkdownToWordConverter:
    """
    Markdown 到 Word 的专业转换器

    相比正则表达式，更稳定地处理 Markdown 语法
    """

    def __init__(self):
        self.lines = []
        self.current_idx = 0

    def convert(self, markdown_content: str, doc) -> None:
        """
        将 Markdown 内容转换并添加到 Word 文档

        Args:
            markdown_content: Markdown 文本
            doc: docx.Document 对象
        """
        self.lines = markdown_content.split('\n')
        self.current_idx = 0

        while self.current_idx < len(self.lines):
            line = self.lines[self.current_idx].rstrip()

            # 空行
            if not line:
                doc.add_paragraph()
                self.current_idx += 1
                continue

            # 标题（# 开头）
            if line.startswith('#'):
                self._add_heading(doc, line)
                self.current_idx += 1
                continue

            # 分隔线（--- 或 ***）
            if line.strip() in ('---', '***', '___'):
                doc.add_paragraph('_' * 50)
                self.current_idx += 1
                continue

            # 代码块（``` 开头）
            if line.strip().startswith('```'):
                self._add_code_block(doc)
                continue

            # 引用（> 开头）
            if line.startswith('>'):
                self._add_blockquote(doc, line)
                self.current_idx += 1
                continue

            # 列表（- 或 * 或 数字. 开头）
            if re.match(r'^(\s*)([-*]|\d+\.)\s+', line):
                self._add_list_item(doc, line)
                self.current_idx += 1
                continue

            # 普通段落（可能跨行）
            self._add_paragraph(doc)

    def _add_heading(self, doc, line: str) -> None:
        """添加标题"""
        level = min(len(line) - len(line.lstrip('#')), 6)
        text = line.lstrip('#').strip()

        heading = doc.add_heading(text, level=level)

        # 处理标题中的行内格式
        self._apply_inline_formatting(heading, text)

    def _add_paragraph(self, doc) -> None:
        """添加段落（可能跨行）"""
        paragraph_lines = []

        while self.current_idx < len(self.lines):
            line = self.lines[self.current_idx].rstrip()

            # 遇到特殊语法，结束段落
            if (not line or
                line.startswith('#') or
                line.strip() in ('---', '***', '___') or
                re.match(r'^(\s*)([-*]|\d+\.)\s+', line) or
                line.strip().startswith('```') or
                line.startswith('>')):
                break

            paragraph_lines.append(line)
            self.current_idx += 1

        if paragraph_lines:
            para_text = "\n".join(paragraph_lines)
            para = doc.add_paragraph()
            self._apply_inline_formatting(para, para_text)

    def _add_list_item(self, doc, line: str) -> None:
        """添加列表项"""
        # 判断列表类型
        if re.match(r'^\s*\d+[\.\)]\s+', line):
            style = 'List Number'
        else:
            style = 'List Bullet'

        # 移除列表标记
        text = re.sub(r'^(\s*)([-*]|\d+[\.\)])\s+', '', line)

        para = doc.add_paragraph(style=style)
        self._apply_inline_formatting(para, text)

    def _add_code_block(self, doc) -> None:
        """添加代码块"""
        # 跳过开始的 ```
        self.current_idx += 1

        code_lines = []
        while self.current_idx < len(self.lines):
            line = self.lines[self.current_idx]
            if line.strip().startswith('```'):
                self.current_idx += 1
                break
            code_lines.append(line)
            self.current_idx += 1

        code_text = "\n".join(code_lines)
        para = doc.add_paragraph(code_text)
        para.style = 'No Spacing'
        try:
            from docx.shared import Pt
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        except ImportError:
            pass

    def _add_blockquote(self, doc, line: str) -> None:
        """添加引用"""
        text = line.lstrip('>').strip()
        para = doc.add_paragraph(text)
        # 引用样式：斜体
        for run in para.runs:
            run.italic = True

    def _apply_inline_formatting(self, element, text: str) -> None:
        """应用行内格式（粗体、斜体、代码）"""
        if hasattr(element, 'clear'):
            element.clear()

        # 定义格式模式：{模式: (前缀长度, 属性设置函数)}
        def _set_bold(run):
            run.bold = True

        def _set_italic(run):
            run.italic = True

        def _set_code(run):
            run.font.name = 'Courier New'

        # 格式映射：模式 -> (前缀长度, 属性设置函数)
        formats = {
            '**': (2, _set_bold),
            '__': (2, _set_bold),
            '_': (1, _set_italic),
            '`': (1, _set_code),
        }

        # 解析行内格式
        parts = re.split(r'(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*|`.*?`)', text)

        for part in parts:
            if not part:
                continue

            # 尝试匹配格式
            matched = False
            for marker, (prefix_len, setter) in formats.items():
                if part.startswith(marker) and part.endswith(marker):
                    # 排除 __ 被 _ 误匹配的情况
                    if marker == '_' and part.startswith('__'):
                        continue
                    content = part[prefix_len:-prefix_len]
                    run = element.add_run(content)
                    setter(run)
                    matched = True
                    break

            if not matched:
                element.add_run(part)


# ========== 原有的辅助函数 ==========

def format_date_for_filename(start_date: str, end_date: str) -> str:
    """
    格式化日期范围用于文件名

    Args:
        start_date: 开始日期（YYYY-MM-DD格式）
        end_date: 结束日期（YYYY-MM-DD格式）

    Returns:
        格式化的日期字符串，如："20250113-20250117"
    """
    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d")

    start_formatted = start_dt.strftime("%Y%m%d")
    end_formatted = end_dt.strftime("%Y%m%d")

    return f"{start_formatted}-{end_formatted}"


def generate_auto_filename(start_date: str, end_date: str) -> str:
    """
    自动生成文件名（默认 Markdown 格式）

    Args:
        start_date: 开始日期（YYYY-MM-DD格式）
        end_date: 结束日期（YYYY-MM-DD格式）

    Returns:
        文件名，如："周报20250113-20250117.md"
    """
    date_str = format_date_for_filename(start_date, end_date)
    return f"周报{date_str}.md"


def detect_output_format(filename: str) -> str:
    """
    根据文件名检测输出格式

    Args:
        filename: 文件名

    Returns:
        输出格式（markdown/word）
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".docx":
        return "word"
    else:
        return "markdown"


def write_markdown_file(content: str, output_path: str, filename: str) -> str:
    """
    写入Markdown文件

    Args:
        content: 完整的周报内容（AI已生成）
        output_path: 输出目录路径
        filename: 文件名

    Returns:
        生成的文件完整路径
    """
    # 确保输出目录存在
    output_dir = Path(output_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 保存文件
    output_file = output_dir / filename
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(content)

    return str(output_file.resolve())


def write_word_file(content: str, output_path: str, filename: str) -> str:
    """
    写入Word文件，使用新的 Markdown → Word 转换器

    Args:
        content: 完整的周报内容（Markdown格式）
        output_path: 输出目录路径
        filename: 文件名

    Returns:
        生成的文件完整路径
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "需要安装 python-docx 库：pip install python-docx\n"
            "或使用 pandoc 转换：pandoc input.md -o output.docx"
        )

    # 确保输出目录存在
    output_dir = Path(output_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    # 创建 Word 文档
    doc = Document()

    # 使用新的专业转换器
    converter = MarkdownToWordConverter()
    converter.convert(content, doc)

    # 保存文件
    output_file = output_dir / filename
    doc.save(output_file)

    return str(output_file.resolve())




def export_report(
    content: str,
    output_path: str,
    filename: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None
) -> str:
    """
    导出周报文件，根据 filename 扩展名自动检测输出格式

    Args:
        content: 完整的周报内容（AI已生成，Markdown格式）
        output_path: 输出目录路径
        filename: 文件名（可选，不提供则自动生成，默认 .md）
        start_date: 开始日期（可选，用于自动文件名）
        end_date: 结束日期（可选，用于自动文件名）

    Returns:
        生成的文件完整路径
    """
    # 如果未提供文件名，自动生成（默认 Markdown 格式）
    if not filename:
        if not start_date or not end_date:
            raise ValueError("未提供文件名时，必须提供 start_date 和 end_date 用于自动生成文件名")
        filename = generate_auto_filename(start_date, end_date)

    # 根据文件名扩展名自动检测输出格式
    output_format = detect_output_format(filename)

    # 根据格式调用对应的写入函数
    if output_format == "markdown":
        return write_markdown_file(content, output_path, filename)
    elif output_format == "word":
        return write_word_file(content, output_path, filename)
    else:
        raise ValueError(f"不支持的文件格式：{Path(filename).suffix}，请使用 .md 或 .docx")


def _export_and_print(
    content: str,
    output_path: str,
    filename: Optional[str],
    start_date: Optional[str],
    end_date: Optional[str]
) -> None:
    """导出周报并打印成功消息"""
    try:
        output_file = export_report(
            content=content,
            output_path=output_path,
            filename=filename,
            start_date=start_date,
            end_date=end_date
        )

        # 检测输出格式
        output_format = detect_output_format(Path(output_file).name)
        format_name = "Word" if output_format == "word" else "Markdown"

        print(f"✅ 已保存 {format_name} 文件到：{output_file}")

    except Exception as e:
        print(f"❌ 导出失败：{e}")
        print("提示：请检查输出目录是否存在，或安装 python-docx：pip install python-docx")
        sys.exit(1)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="将AI生成的周报内容写入文件，支持Markdown/Word格式，支持自定义模板"
    )
    parser.add_argument(
        "--content",
        type=str,
        help="完整的周报内容（AI已生成，Markdown格式）"
    )
    parser.add_argument(
        "--data",
        type=str,
        help="JSON 数据文件路径（用于填充 Word 模板）"
    )
    parser.add_argument(
        "--template",
        type=str,
        help="模板文件路径（支持 .md、.docx）"
    )
    parser.add_argument(
        "--output",
        type=str,
        required=True,
        help="输出目录路径"
    )
    parser.add_argument(
        "--filename",
        type=str,
        help="输出文件名（可选，不提供则自动生成）"
    )
    parser.add_argument(
        "--start-date",
        type=str,
        help="开始日期（YYYY-MM-DD格式，用于自动文件名）"
    )
    parser.add_argument(
        "--end-date",
        type=str,
        help="结束日期（YYYY-MM-DD格式，用于自动文件名）"
    )

    args = parser.parse_args()

    # 检查是否提供了模板
    if args.template:
        # 模板模式
        template_path = Path(args.template)

        if not template_path.exists():
            print(f"❌ 模板文件不存在：{args.template}")
            sys.exit(1)

        # 根据模板类型处理
        if template_path.suffix.lower() == ".docx":
            # Word 模板：使用 fill_template.py 填充
            if not args.data:
                print("❌ Word 模板需要提供 --data 参数（JSON 数据文件）")
                sys.exit(1)

            if not args.filename:
                print("❌ Word 模板需要提供 --filename 参数")
                sys.exit(1)

            try:
                # 导入 fill_template 模块
                import importlib.util
                fill_template_path = Path(__file__).parent / "fill_template.py"
                spec = importlib.util.spec_from_file_location("fill_template", fill_template_path)
                fill_template_module = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(fill_template_module)

                # 调用填充函数
                output_file = fill_template_module.fill_word_template(
                    template_path=args.template,
                    data_path=args.data,
                    output_path=args.output,
                    filename=args.filename
                )

                print(f"✅ 已填充 Word 模板并保存到：{output_file}")

            except Exception as e:
                print(f"❌ 填充 Word 模板失败：{e}")
                print("提示：请检查模板文件和数据文件是否正确")
                sys.exit(1)

        elif template_path.suffix.lower() == ".md":
            # Markdown 模板：直接处理内容
            if not args.content:
                print("❌ Markdown 模板需要提供 --content 参数")
                sys.exit(1)
            _export_and_print(args.content, args.output, args.filename, args.start_date, args.end_date)

        else:
            print(f"❌ 不支持的模板类型：{template_path.suffix}")
            print("💡 支持的模板类型：.md、.docx")
            sys.exit(1)

    else:
        # 原有模式：直接导出
        if not args.content:
            print("❌ 未提供模板时，必须提供 --content 参数")
            sys.exit(1)
        _export_and_print(args.content, args.output, args.filename, args.start_date, args.end_date)


if __name__ == "__main__":
    main()
